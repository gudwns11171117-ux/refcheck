# -*- coding: utf-8 -*-
"""문서(PDF/DOCX/HWPX/TXT)에서 텍스트를 뽑고, 참고문헌 구역을 찾아 항목 단위로 나눈다."""
from __future__ import annotations

import io
import re
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from typing import Optional

import pymupdf


@dataclass
class Line:
    text: str
    x0: float = 0.0
    y0: float = 0.0
    x1: float = 0.0
    page: int = 0
    col: int = 0
    size: float = 0.0


@dataclass
class ExtractResult:
    refs: list[str]
    section_text: str
    heading_found: bool
    heading_text: str
    method: str          # numbered | indent | heuristic | paragraph
    pages: int
    warnings: list[str]


# ---------------------------------------------------------------- 헤딩/종료 패턴
_KW = (
    r"참\s*고\s*문\s*헌(?:\s*목\s*록)?|인\s*용\s*문\s*헌|참\s*고\s*자\s*료|참\s*고\s*서\s*지|"
    r"References?|REFERENCES?|Reference\s+List|Bibliography|BIBLIOGRAPHY|"
    r"Literature\s+Cited|Works\s+Cited|參考文獻|参考文献"
)
HEADING_RE = re.compile(
    r"^\s*(?:[\[\(<【]?\s*(?:[IVXⅠ-Ⅻ]{1,4}|\d{1,2}|[가-힣])\s*[\]\)>】\.]?\s*)?"
    r"(?:■|▣|◆|●|□|○|▶|▷|※)?\s*"
    r"(" + _KW + r")"
    r"(?:\s*[\(（]\s*(?:" + _KW + r")\s*[\)）])?"
    r"\s*[:：]?\s*$",
    re.I,
)
END_RE = re.compile(
    r"^\s*(?:[\[\(<【〈]?\s*)?(?:Abstract|ABSTRACT|국\s*문\s*초\s*록|영\s*문\s*초\s*록|초\s*록|Appendix|APPENDIX|"
    r"부\s*록|저\s*자\s*소\s*개|저자\s*약력|Author\s+Biograph\w*|About\s+the\s+Authors?|"
    r"Acknowledg\w*|감사의\s*글|Supplementary|Correspondence|논문\s*접수|접수일|투고일|"
    r"Received[:\s]|Summary|要\s*約|摘\s*要)",
    re.I,
)
PAGE_NUM_RE = re.compile(r"^\s*[-–—─]?\s*\d{1,4}\s*[-–—─]?\s*$")
NUM_RE = re.compile(r"^\s*(?:\[\s*(\d{1,3})\s*\]|\(\s*(\d{1,3})\s*\)|(\d{1,3})\s*[\.\)]|(\d{1,3})\s*〕)\s*")
YEAR_RE = re.compile(r"(?<!\d)(?:1[89]\d{2}|20\d{2})(?!\d)")

# 새 항목의 첫 줄처럼 보이는 패턴(저자명 시작)
AUTHOR_START = re.compile(
    r"^(?:"
    r"[가-힣]{2,5}(?:\s*[,·・]\s*|\s*[\(（]|\s+외|\s+등|\s*\.|\s*$)"      # 홍길동, / 홍길동( / 홍길동 외
    r"|[A-Z][A-Za-z'’\-]+,\s*(?:[A-Z]\.|[A-Z][a-z]+)"                    # Kim, J. / Kim, John
    r"|[A-Z]\.\s*[A-Z]?\.?\s*[A-Z][a-z]+"                                # J. K. Kim
    r"|[A-Z][a-z]+\s+[A-Z]\.\s*[A-Z]?"                                   # John K.
    r"|[A-Z][A-Za-z'’\-]+\s+et\s+al"                                     # Kim et al
    r"|[A-Z]{2,}(?:\s|,|\.)"                                             # KOSHA, / OECD
    r"|(?:[A-Z][a-z]+\s+){1,5}(?:of|for|and|on)\s"                       # Ministry of ...
    r"|[「『\"“‘]"                                                      # 따옴표로 시작
    r"|고용노동부|산업안전보건공단|안전보건공단|국토교통부|보건복지부|교육부|환경부|통계청|행정안전부|"
    r"한국[가-힣]{2,10}(?:원|회|단|부|청|소)"
    r")"
)
ENDS_TERMINAL = re.compile(r"(?:[\.\)\]」』\"”’]|\d)\s*$")
EARLY_YEAR = re.compile(r"^.{0,60}?[\(（]\s*(?:1[89]\d{2}|20\d{2})[a-z]?\s*[\)）]")


# ---------------------------------------------------------------- 문서 → 줄
def _pdf_lines(data: bytes) -> tuple[list[Line], int]:
    doc = pymupdf.open(stream=data, filetype="pdf")
    all_lines: list[Line] = []
    for pno, page in enumerate(doc):
        width = page.rect.width or 1.0
        d = page.get_text("dict")
        plines: list[Line] = []
        for b in d.get("blocks", []):
            if b.get("type") != 0:
                continue
            for ln in b.get("lines", []):
                txt = "".join(s.get("text", "") for s in ln.get("spans", []))
                if not txt.strip():
                    continue
                x0, y0, x1, y1 = ln["bbox"]
                size = max((s.get("size", 0) for s in ln.get("spans", [])), default=0.0)
                plines.append(Line(txt, x0, y0, x1, pno, 0, size))
        if not plines:
            continue
        # 2단 판형 감지: 줄 너비가 페이지 절반 이하인 줄이 대다수면 2단
        narrow = sum(1 for l in plines if (l.x1 - l.x0) < 0.55 * width)
        two_col = narrow >= 0.6 * len(plines) and len(plines) >= 8
        if two_col:
            for l in plines:
                center = (l.x0 + l.x1) / 2
                wide = (l.x1 - l.x0) > 0.6 * width
                l.col = 0 if (wide or center < width / 2) else 1
            plines.sort(key=lambda l: (l.col, round(l.y0, 1), l.x0))
        else:
            plines.sort(key=lambda l: (round(l.y0, 1), l.x0))
        all_lines.extend(plines)
    n_pages = len(doc)
    doc.close()
    return all_lines, n_pages


def _docx_lines(data: bytes) -> list[Line]:
    import docx  # python-docx
    d = docx.Document(io.BytesIO(data))
    out = []
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t:
            out.append(Line(t, 0, float(i), 0, 0))
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    out.append(Line(t, 0, 0, 0, 0))
    return out


def _hwpx_lines(data: bytes) -> list[Line]:
    out = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        names = sorted(n for n in z.namelist() if re.match(r"Contents/section\d+\.xml$", n))
        for n in names:
            xml = z.read(n).decode("utf-8", errors="replace")
            for para in re.findall(r"<hp:p\b.*?</hp:p>", xml, flags=re.S):
                texts = re.findall(r"<hp:t[^>]*>(.*?)</hp:t>", para, flags=re.S)
                t = "".join(texts)
                t = re.sub(r"<[^>]+>", "", t)
                t = (t.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
                     .replace("&quot;", '"').replace("&apos;", "'"))
                t = t.strip()
                if t:
                    out.append(Line(t, 0, float(len(out)), 0, 0))
    return out


def _txt_lines(data: bytes) -> list[Line]:
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr", "utf-16"):
        try:
            s = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        s = data.decode("utf-8", errors="replace")
    return [Line(t.rstrip(), 0, float(i), 0, 0) for i, t in enumerate(s.splitlines())]


# ---------------------------------------------------------------- 머리말/꼬리말 제거
def _drop_running_lines(lines: list[Line]) -> list[Line]:
    pages_of: dict[str, set[int]] = defaultdict(set)
    for l in lines:
        key = re.sub(r"\d+", "#", l.text.strip())
        if len(key) <= 80:
            pages_of[key].add(l.page)
    out = []
    for l in lines:
        t = l.text.strip()
        key = re.sub(r"\d+", "#", t)
        if PAGE_NUM_RE.match(t):
            continue
        if len(pages_of.get(key, ())) >= 3 and len(t) <= 80:
            continue
        out.append(l)
    return out


# ---------------------------------------------------------------- 참고문헌 구역 찾기
def find_reference_section(lines: list[Line]) -> tuple[Optional[list[Line]], str]:
    cands = [i for i, l in enumerate(lines) if len(l.text.strip()) <= 40 and HEADING_RE.match(l.text.strip())]
    start = None
    heading = ""
    for i in reversed(cands):
        if len(lines) - (i + 1) >= 3:
            start = i + 1
            heading = lines[i].text.strip()
            break
    if start is None:
        # 헤딩과 첫 항목이 한 줄에 붙은 경우: "참고문헌 홍길동(2020)..."
        inline = re.compile(r"^\s*(?:\d{1,2}\.|[IVXⅠ-Ⅻ]+\.)?\s*(참고문헌|References?|REFERENCES)\b[\s:：]*(.+)$", re.I)
        for i in range(len(lines) - 1, -1, -1):
            m = inline.match(lines[i].text.strip())
            if m and len(m.group(2)) > 15:
                new = Line(m.group(2), lines[i].x0, lines[i].y0, lines[i].x1, lines[i].page, lines[i].col, lines[i].size)
                lines = lines[:i] + [new] + lines[i + 1:]
                start = i
                heading = m.group(1)
                break
    if start is None:
        return None, ""
    end = len(lines)
    for j in range(start, len(lines)):
        t = lines[j].text.strip()
        if len(t) <= 40 and END_RE.match(t):
            end = j
            break
    return lines[start:end], heading


# ---------------------------------------------------------------- 항목 나누기
def _longest_chain(markers: list[tuple[int, int]]) -> list[tuple[int, int]]:
    """(줄번호, 번호) 목록에서 1씩 증가하는 가장 긴 사슬(최대 1개 건너뜀 허용)."""
    best: list[tuple[int, int]] = []
    for s in range(len(markers)):
        if markers[s][1] > 3:      # 1~3 근처에서 시작하는 사슬만
            continue
        chain = [markers[s]]
        for k in range(s + 1, len(markers)):
            if markers[k][1] in (chain[-1][1] + 1, chain[-1][1] + 2):
                chain.append(markers[k])
        if len(chain) > len(best):
            best = chain
    return best


def _join_group(texts: list[str]) -> str:
    out = ""
    for t in texts:
        t = t.strip()
        if not t:
            continue
        if not out:
            out = t
            continue
        if out.endswith("-") and t[:1].islower():
            out = out[:-1] + t
        elif out.endswith(("-", "–")) and re.match(r"^\d", t):
            out = out + t
        else:
            out = out + " " + t
    out = re.sub(r"\s+", " ", out).strip()
    return out


def _join_groups(texts: list[str], starts: list[int], strip_marker: bool = False) -> list[str]:
    starts = sorted(set(starts))
    if not starts or starts[0] != 0:
        starts = [0] + starts
    refs = []
    for a, b in zip(starts, starts[1:] + [len(texts)]):
        chunk = texts[a:b]
        if strip_marker and chunk:
            chunk = [NUM_RE.sub("", chunk[0], count=1)] + chunk[1:]
        s = _join_group(chunk)
        if s:
            refs.append(s)
    return refs


def _indent_starts(section: list[Line]) -> Optional[list[int]]:
    groups: dict[tuple[int, int], list[int]] = defaultdict(list)
    for i, l in enumerate(section):
        groups[(l.page, l.col)].append(i)
    informative = 0
    starts: list[int] = []
    ambiguous_groups: list[list[int]] = []
    for key, idxs in groups.items():
        xs = [section[i].x0 for i in idxs]
        xl = min(xs)
        firsts = [i for i in idxs if section[i].x0 <= xl + 2.5]
        conts = [i for i in idxs if section[i].x0 > xl + 4.0]
        if firsts and conts:
            informative += 1
            starts.extend(firsts)
        else:
            ambiguous_groups.append(idxs)
    if informative == 0:
        return None
    # 들여쓰기 정보가 없는 그룹은 텍스트 휴리스틱으로 보완
    texts = [l.text for l in section]
    for idxs in ambiguous_groups:
        sub = [texts[i] for i in idxs]
        for k in _heuristic_starts(sub):
            starts.append(idxs[k])
    return sorted(set(starts))


def _heuristic_starts(texts: list[str]) -> list[int]:
    starts = [0]
    cur: list[str] = []
    for i, t in enumerate(texts):
        t = t.strip()
        if i == 0:
            cur = [t]
            continue
        prev = texts[i - 1].strip()
        cur_text = " ".join(cur)
        cur_ok = bool(YEAR_RE.search(cur_text)) or len(cur_text) >= 40
        is_start = AUTHOR_START.match(t) and (ENDS_TERMINAL.search(prev) or EARLY_YEAR.match(t)) and cur_ok
        if is_start:
            starts.append(i)
            cur = [t]
        else:
            cur.append(t)
    return starts


def _paragraph_refs(texts: list[str]) -> list[str]:
    """DOCX/HWPX처럼 문단 단위가 곧 항목인 경우. 너무 짧은 조각은 앞 항목에 붙인다."""
    refs: list[str] = []
    for t in texts:
        t = re.sub(r"\s+", " ", t).strip()
        if not t:
            continue
        if refs and (len(t) < 25 or (not YEAR_RE.search(t) and not AUTHOR_START.match(t))) and not NUM_RE.match(t):
            refs[-1] = _join_group([refs[-1], t])
        else:
            refs.append(t)
    return [NUM_RE.sub("", r, count=1).strip() if NUM_RE.match(r) else r for r in refs]


def split_references(section: list[Line], has_geom: bool, para_mode: bool) -> tuple[list[str], str]:
    texts = [l.text for l in section]
    # 1) 번호 붙은 목록
    markers = []
    for i, t in enumerate(texts):
        m = NUM_RE.match(t)
        if m:
            n = int(next(g for g in m.groups() if g))
            markers.append((i, n))
    chain = _longest_chain(markers)
    if len(chain) >= 3 and len(chain) >= 0.5 * chain[-1][1]:
        starts = [i for i, _ in chain]
        lead = texts[:starts[0]]
        if lead and len(" ".join(lead)) < 20:      # 헤딩 잔여물
            texts = texts[starts[0]:]
            starts = [i - chain[0][0] for i, _ in chain]
        return _join_groups(texts, starts, strip_marker=True), "numbered"
    # 2) 문단 단위 문서
    if para_mode:
        return _paragraph_refs(texts), "paragraph"
    # 3) 들여쓰기(PDF 좌표)
    if has_geom:
        st = _indent_starts(section)
        if st and len(st) >= 2:
            refs = _join_groups(texts, st)
            bad = sum(1 for r in refs if not YEAR_RE.search(r))
            if bad <= len(refs) * 0.5:
                return refs, "indent"
    # 4) 텍스트 휴리스틱
    return _join_groups(texts, _heuristic_starts(texts)), "heuristic"


# ---------------------------------------------------------------- 진입점
def extract_references(data: bytes, filename: str) -> ExtractResult:
    name = (filename or "").lower()
    warnings: list[str] = []
    has_geom = False
    para_mode = False
    pages = 0
    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        lines, pages = _pdf_lines(data)
        has_geom = True
        if not lines:
            warnings.append("PDF에서 텍스트를 추출하지 못했습니다. 스캔 이미지 PDF라면 OCR이 필요합니다.")
        lines = _drop_running_lines(lines)
    elif name.endswith(".docx"):
        lines = _docx_lines(data)
        para_mode = True
    elif name.endswith(".hwpx"):
        lines = _hwpx_lines(data)
        para_mode = True
    elif name.endswith(".hwp"):
        raise ValueError("구버전 .hwp 파일은 지원하지 않습니다. 한글에서 '.hwpx' 또는 'PDF'로 저장해 주세요.")
    else:
        lines = _txt_lines(data)
        para_mode = True

    section, heading = find_reference_section(lines)
    heading_found = section is not None
    if section is None:
        warnings.append("'참고문헌/References' 제목을 찾지 못해 문서 뒷부분(30%)을 대상으로 했습니다. 목록을 확인해 주세요.")
        cut = int(len(lines) * 0.7)
        section = lines[cut:] if lines else []
    refs, method = split_references(section, has_geom, para_mode)
    refs = [r for r in refs if len(r) >= 12]
    section_text = "\n".join(l.text for l in section)
    return ExtractResult(refs, section_text, heading_found, heading, method, pages, warnings)


def split_pasted_text(text: str) -> list[str]:
    """사용자가 붙여넣은 참고문헌 텍스트를 항목 단위로 나눈다."""
    raw_lines = [l.rstrip() for l in text.splitlines()]
    # 빈 줄로 구분된 블록이 있으면 블록 = 항목
    blocks: list[list[str]] = [[]]
    for l in raw_lines:
        if not l.strip():
            if blocks[-1]:
                blocks.append([])
        else:
            blocks[-1].append(l)
    blocks = [b for b in blocks if b]
    if len(blocks) >= 3:
        joined = [_join_group(b) for b in blocks]
        if sum(1 for j in joined if YEAR_RE.search(j)) >= 0.5 * len(joined):
            return [NUM_RE.sub("", j, count=1).strip() if NUM_RE.match(j) else j for j in joined if len(j) >= 12]
    lines = [Line(l, 0, float(i), 0, 0) for i, l in enumerate(raw_lines) if l.strip()]
    texts = [l.text for l in lines]
    with_year = sum(1 for t in texts if YEAR_RE.search(t))
    if texts and with_year >= 0.6 * len(texts):
        refs, _ = split_references(lines, False, True)
    else:
        refs, _ = split_references(lines, False, False)
    return [r for r in refs if len(r) >= 12]
